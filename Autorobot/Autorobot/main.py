import re
import tkinter
import tkinter.messagebox
from tkinter.messagebox import showinfo

from threading import Thread
from time import sleep
import customtkinter as ctk
import json
import os
from datetime import datetime
from cryptography.fernet import Fernet

import ctypes, sys

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from dotenv import load_dotenv
import requests

from update import Update

current_dir = os.path.dirname(__file__)

load_dotenv(dotenv_path=os.path.join(current_dir, ".env"))

encryption_key = os.getenv("ENCRYPTION_KEY")
cipher_suite = Fernet(encryption_key.encode())


# ctk.set_appearance_mode("dark")
# ctk.set_default_color_theme("dark-blue")

INTERNET_CONNECTED = False


class DownloadDialog(ctk.CTk):
    def __init__(self, master: ctk.CTkToplevel):
        super().__init__()

        self.Main = master

        self.title("Download Update")
        self.geometry(f"{400}x{100}")
        self.resizable(width=False, height=False)
        self.protocol("WM_DELETE_WINDOW", self.on_closing)

        # self.slider_progressbar_frame = ctk.CTkFrame(self, fg_color="transparent")
        # self.slider_progressbar_frame.grid(row=0, column=0, padx=20, pady=20, sticky="nsw")

        self.progressbar_1 = ctk.CTkProgressBar(self, width=300, height=20, corner_radius=8)
        self.progressbar_1.pack(pady="10")
        self.progressbar_1.set(0)

        self.percent = ctk.StringVar()
        self.percent_label = ctk.CTkLabel(self, text=self.percent.get())
        self.percent_label.pack()

        self.button_download = ctk.CTkButton(master=self, text="Download", command=self.Start)
        self.button_download.pack()
        # self.progressbar_1.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")

    def Start(self):
        global INTERNET_CONNECTED

        if INTERNET_CONNECTED:
            self.patch = Update()
            self.button_download.configure(state="disabled")
            thread_updater = Thread(
                target=self.patch.ApplyUpdate, kwargs={"DownloadDialog": self, "Master": self.Main}
            )
            thread_updater.daemon = True
            thread_updater.start()
        else:
            showinfo("Autorobot", "Lütfen İnternet Bağlantınızı Kontrol Edip Tekrar Deneyin")

    def on_closing(self):
        self.destroy()


class MasterGui(ctk.CTk):
    def __init__(self):
        global current_dir
        super().__init__()

        # create tabview
        self.tabview = ctk.CTkTabview(self, height=300, width=300)
        self.tabview.grid(row=0, column=1, padx=(20, 20), pady=(0, 0), sticky="nsew")
        self.tabview.add("Ayarlar")
        self.tabview.add("Hakkında")
        self.tabview.tab("Ayarlar").grid_columnconfigure(
            0, weight=1
        )  # configure grid of individual tabs
        self.tabview.tab("Hakkında").grid_columnconfigure(0, weight=1)
        self.texbox_about = ctk.CTkTextbox(
            master=self.tabview.tab("Hakkında"), width=250, height=250
        )
        self.texbox_about.grid(row=0, column=0)
        self.texbox_about.insert(
            ctk.END,
            """
        ------------------------------------------------------
            Autorobot v0.2.0 Hakkında      
        ------------------------------------------------------

Autorobot ile kolay docx formları oluşturun ve zaman kazanın.

Version: 0.2.0
Release Date: November 15, 2023

**Credits:**
- Developer: Kılıçarslan SIMSIKI
- UI/UX Design: Kılıçarslan SIMSIKI

**Contact and Support:**
For any inquiries, feedback, or technical support,\n please contact our dedicated support team at https://swordlion.org

**Legal Information:**
Autorobot is a trademark of swordlion.org Corporation. All rights reserved.
This software is protected by international copyright laws.
Unauthorized reproduction or distribution of this software is prohibited.

**Privacy Policy:**
Read our privacy policy at https://swordlion.org for\n information on how we handle your data and protect your privacy.

Thank you for choosing Autorobot.


            """,
        )
        self.texbox_about.configure(state="disabled")
        self.button_update = ctk.CTkButton(
            master=self.tabview.tab("Ayarlar"),
            command=self.Update_Control,
            text="Güncellemeyi İndir",
        )
        self.button_update.grid(row=0, column=0, padx=20, pady=(10, 10))

        self.total_cost = 0
        self.form_date_paragraph = None

        # Create or open the Word document
        self.desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        self.document_path = os.path.join(
            self.desktop_path, cipher_suite.decrypt(os.getenv("DOCUMENT_PATH").encode()).decode()
        )
        self.document = Document()
        if os.path.exists(self.document_path):
            self.document = Document(self.document_path)

        # Initialize total cost by scanning existing user entries
        self.total_cost = self.calculate_total_cost()

        # Call Update_Form_Date in __init__
        self.Update_Form_Date()

        # Gui setup
        self.title("Autorobot")
        self.geometry(f"{880}x{350}")
        self.resizable(width=False, height=False)
        self.protocol("WM_DELETE_WINDOW", self.on_closing)

        self.after(201, lambda: self.iconbitmap(os.path.join(current_dir, "ar.ico")))

        self.grid_rowconfigure((0, 1, 2, 3), weight=1)
        self.grid_columnconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=1)

        # Frame 1 and Border Frame
        self.border_frame = ctk.CTkFrame(master=self)
        self.border_frame.grid(row=0, column=0, padx=20, pady=20, sticky="nsw")

        self.frame_1 = ctk.CTkFrame(master=self.border_frame)
        self.frame_1.grid(row=0, column=0, padx=20, pady=20, sticky="ns")

        self.DateEntry_label = ctk.CTkLabel(
            master=self.frame_1, text="Tarih", font=("Segoe UI", 18)
        )
        self.DateEntry_label.grid(row=0, column=0, sticky="nsew")
        self.DateEntry = ctk.CTkEntry(master=self.frame_1, placeholder_text="01.01.2023")
        self.DateEntry.grid(row=1, column=0, padx=10, pady=10, sticky="nsew")

        self.NoteEntry_label = ctk.CTkLabel(
            master=self.frame_1, text="Açıklama", font=("Segoe UI", 18)
        )
        self.NoteEntry_label.grid(row=0, column=1, sticky="nsew")
        self.NoteEntry = ctk.CTkEntry(master=self.frame_1, placeholder_text="Açıklama girin...")
        self.NoteEntry.grid(row=1, column=1, padx=10, pady=10, sticky="nsew")

        self.CostEntry_label = ctk.CTkLabel(
            master=self.frame_1, text="Tutar", font=("Segoe UI", 18)
        )
        self.CostEntry_label.grid(row=0, column=2, sticky="nsew")
        self.CostEntry = ctk.CTkEntry(master=self.frame_1, placeholder_text="140")
        self.CostEntry.grid(row=1, column=2, padx=10, pady=10, sticky="nsew")

        self.button_1 = ctk.CTkButton(
            master=self.frame_1,
            command=self.Save_Document,
            text="Kaydet",
        )
        self.button_1.grid(row=2, column=0, columnspan=3, pady=10, padx=10, sticky="ns")

        self.isUpdateAvailable = False

    def on_closing(self):
        """Called when you press the X button to close the program. Kills the GUI and the opened chromedriver threads"""
        # Update the total row in the table
        # self.download_dialog.destroy()
        self.destroy()
        sys.exit()

    def start_internet_check_thread(self):
        # Create a separate thread for checking internet connection
        internet_check_thread = Thread(target=self.check_internet_connection)
        # Set the thread as a daemon so that it will exit when the main program exits
        internet_check_thread.daemon = True
        # Start the thread
        internet_check_thread.start()

    def check_internet_connection(self):
        global INTERNET_CONNECTED
        print(INTERNET_CONNECTED)
        while True:
            try:
                # Try making a simple request to a known server
                response = requests.get("https://www.google.com", timeout=5)
                # If the request is successful, the internet is connected
                INTERNET_CONNECTED = True
            except requests.exceptions.RequestException:
                # If an exception occurs, the internet is not connected
                INTERNET_CONNECTED = False

            # Wait for a certain interval before checking again
            print(INTERNET_CONNECTED)
            sleep(5)  # Adjust the interval as needed

    def Update_Init(self):
        global INTERNET_CONNECTED
        self.patch = Update()

        if INTERNET_CONNECTED:
            self.isUpdateAvailable = self.patch.IsUpdateAvailable()

        if self.isUpdateAvailable:
            self.Show_Update_Popup()
        else:
            pass

    def Show_Update_Popup(self):
        showinfo(
            "Autorobot",
            "Yeni Güncelleme Mevcut!\nAyarlar->Güncelleme butonundan güncellemeyi indirin",
        )

    def Update_Control(self):
        global INTERNET_CONNECTED
        if INTERNET_CONNECTED:
            if self.isUpdateAvailable:
                self.download_dialog = DownloadDialog(master=self)
                self.download_dialog.mainloop()
            else:
                showinfo("Autorobot", "Yeni Güncelleme Mevcut Değil! İyi Günler :)")
        else:
            showinfo("Autorobot", "Lütfen İnternet Bağlantınızı Kontrol Edip Tekrar Deneyin")

    def Save_Document(self):
        # Get user inputs
        date = self.DateEntry.get()
        note = self.NoteEntry.get()
        cost = self.CostEntry.get()

        # Check if the tables exist, otherwise, create new ones
        if not self.document.tables:
            self.create_user_entry_table()

        # Find the user entries table
        user_entries_table = self.document.tables[0]

        # Add user inputs to a new row in the user entries table
        row = user_entries_table.add_row()
        row.cells[0].text = date
        row.cells[1].text = note
        row.cells[2].text = cost

        # Update the running total
        self.total_cost += float(cost)

        # Find the total cost table and update the total cost value
        if len(self.document.tables) > 1:
            total_cost_table = self.document.tables[1]
            total_cost_table.rows[0].cells[0].text = f"Toplam Tutar: {self.total_cost:.2f}"
            total_cost_table.rows[0].cells[0].paragraphs[0].runs[
                0
            ].bold = True  # Make the text bold
        else:
            # If the total cost table doesn't exist, create it
            self.create_total_cost_table()

        # Sort the user entries table based on the date column
        self.sort_user_entries_table()

        # Save the document
        self.document.save(self.document_path)

    def Update_Form_Date(self):
        import calendar
        from datetime import datetime

        # Recreate the header to get the reference
        self.form_date_paragraph = self.Create_Header()

        # Get the current date
        current_date = datetime.now()
        # Get the last day of the current month
        last_day = calendar.monthrange(current_date.year, current_date.month)[1]
        # Create a new date with the last day of the current month
        last_day_of_month = datetime(current_date.year, current_date.month, last_day)

        # Update the text content of the form date paragraph
        self.form_date_paragraph.text = f"Form Tarihi: {last_day_of_month.strftime('%d.%m.%Y')}"
        # Save the document
        self.document.save(self.document_path)

    def Create_Header(self):
        # Check if the header already exists
        if (
            not self.document.paragraphs
            or "MASRAF BEYAN FORMU" not in self.document.paragraphs[0].text
        ):
            # Add a centered and bold title
            title = self.document.add_heading("MASRAF BEYAN FORMU", level=1)
            self.document.add_paragraph("\n")
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = title.runs[0]
            run.bold = True
            run.font.size = Pt(30)

            # Add personnel name
            self.document.add_paragraph("Personel Adı: xx")

        # Check if the form date paragraph already exists
        if not self.form_date_paragraph:
            # Try to find the form date paragraph in the existing paragraphs
            for paragraph in self.document.paragraphs:
                if "Form Tarihi:" in paragraph.text:
                    self.form_date_paragraph = paragraph
                    break
            # If not found, add the form date paragraph
            if not self.form_date_paragraph:
                self.form_date_paragraph = self.document.add_paragraph()

        return self.form_date_paragraph

    def create_user_entry_table(self):
        # Add a user entries table
        user_entries_table = self.document.add_table(rows=1, cols=3)
        user_entries_table.style = "Table Grid"  # Add grid style to the table
        user_entries_table.autofit = False  # Disable autofit to keep fixed column width

        # Set column widths
        for col in user_entries_table.columns:
            col.width = Pt(100)

        headers = user_entries_table.rows[0].cells
        headers[0].text = "Tarih"
        headers[1].text = "Açıklama"
        headers[2].text = "Tutar"

        # Make headers bold
        for cell in headers:
            cell.paragraphs[0].runs[0].bold = True

        # Add spacing between the tables
        self.document.add_paragraph("\n")

    def sort_user_entries_table(self):
        # Get the user entries table
        user_entries_table = self.document.tables[0]

        # Get the user entries and sort them based on the date column (assuming the date is in the first column)
        user_entries = [
            (row.cells[0].text, row.cells[1].text, row.cells[2].text)
            for row in user_entries_table.rows[1:]
        ]
        user_entries.sort(key=lambda entry: datetime.strptime(entry[0], "%d.%m.%Y"))

        # Clear existing rows in the table
        for _ in range(1, len(user_entries_table.rows)):
            user_entries_table._tbl.remove(user_entries_table.rows[1]._element)

        # Add the sorted rows back to the table
        for entry in user_entries:
            cells = user_entries_table.add_row().cells
            cells[0].text = entry[0]
            cells[1].text = entry[1]
            cells[2].text = entry[2]

    def create_total_cost_table(self):
        # Add a total cost table
        total_cost_table = self.document.add_table(rows=1, cols=1)
        total_cost_table.style = "Table Grid"  # Add grid style to the table
        total_cost_table.autofit = False  # Disable autofit to keep fixed column width

        # Set column width
        total_cost_table.columns[0].width = Pt(300)

        # Add a header to the total cost table
        total_cost_table.rows[0].cells[0].text = "Toplam Tutar: 0.00"

    def calculate_total_cost(self):
        total_cost = 0

        # Check if the tables exist
        if self.document.tables:
            user_entries_table = self.document.tables[0]

            # Iterate over rows starting from the second row (skip headers)
            for row in user_entries_table.rows[1:]:
                cost_cell = row.cells[2]
                try:
                    total_cost += float(cost_cell.text)
                except ValueError:
                    pass  # Handle invalid entries

        return total_cost

    def dummy2(self):
        pass

    def dummy3(self):
        pass


def is_admin():
    try:
        return ctypes.windll.shell32.IsUserAnAdmin()
    except:
        return False


if __name__ == "__main__":
    if is_admin():
        # initate the gui
        Autorobot = MasterGui()
        Autorobot.start_internet_check_thread()
        sleep(2)
        Autorobot.Update_Init()
        # start the gui
        Autorobot.mainloop()
    else:
        # Re-run the program with admin rights
        ctypes.windll.shell32.ShellExecuteW(None, "runas", sys.executable, __file__, None, 1)
