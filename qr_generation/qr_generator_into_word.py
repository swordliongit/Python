import qrcode
from docx import Document
from docx.shared import Inches

# Create a new Word document
document = Document()

# Generate and insert QR codes with labels
for i in range(1000, 1101):
    content = f"XSARJ-{i}"

    # Generate QR code
    qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_L, box_size=10, border=4)
    qr.add_data(content)
    qr.make(fit=True)
    qr_img = qr.make_image(fill_color="black", back_color="white")

    # Save the QR code as an image file
    qr_filename = f"qrcode_{i}.png"
    qr_img.save(qr_filename)
    print(f"Generated QR code for {content}")

    # Insert the QR code into the Word document with the label
    document.add_paragraph(f"XSARJ-{i}")
    document.add_picture(qr_filename, width=Inches(2.0))
    document.add_paragraph("\n")

    # Delete the QR code image file
    import os
    os.remove(qr_filename)

# Save the Word document
document.save("qrcodes.docx")
print("QR codes saved in the Word document.")
